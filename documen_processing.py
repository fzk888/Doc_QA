import os
import re
import pdfplumber
import pandas as pd
from tqdm import tqdm
from langchain.text_splitter import MarkdownHeaderTextSplitter, RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader
from langchain.schema import Document
import pypandoc
import logging
from pypdf import PdfReader
import requests
import yaml

from pdf2image import convert_from_path
import requests
import os
import numpy as np
from PIL import Image
import io
import json
import zipfile

from docx import Document as Document2

url_f = None

# 统一分块日志记录器
chunk_logger = logging.getLogger("docqa.chunk")
with open("config.yaml", "r") as _cf:
    _cfg = yaml.safe_load(_cf)
ENABLE_OCR_IMAGES = bool(_cfg.get("settings", {}).get("enable_ocr_images", False))
ENABLE_PDF_PIX2TEXT = bool(_cfg.get("settings", {}).get("enable_pdf_pix2text", False))
URL_OCR = _cfg.get("paths", {}).get("ocr_service_url", "http://127.0.0.1:8001/detection_pic")
URL_PIX2TEXT = _cfg.get("paths", {}).get("pix2text_url", "http://127.0.0.1:8503/pix2text")
OCR_TIMEOUT = int(_cfg.get("settings", {}).get("ocr_timeout_secs", 30))
url_f = URL_OCR

# Ensure pandoc is available for pypandoc on Windows
def _ensure_pandoc_installed():
    try:
        # Will raise OSError if pandoc is not found
        pypandoc.get_pandoc_version()
    except OSError:
        logging.info("Pandoc not found. Will use python-docx as fallback.")

_ensure_pandoc_installed()

def is_valid_table(table_df, min_char_count=120):
    total_char_count = sum(len(str(cell)) for _, row in table_df.iterrows() for cell in row)
    if total_char_count < min_char_count:
        return False
    return True
def extract_all_table_and_textand_image(pdf_path, out_path_md,image_dir):
    pdf = pdfplumber.open(pdf_path)

    def is_sentence_end(text):
        return re.search(r'[。？！.?!]$', text) is not None

    def not_within_bboxes(obj):
        return not any(
            obj["x0"] + obj["x1"] >= 2 * bbox[0] and obj["x0"] + obj["x1"] < 2 * bbox[2] and
            obj["top"] + obj["bottom"] >= 2 * bbox[1] and obj["top"] + obj["bottom"] < 2 * bbox[3]
            for bbox in bboxes
        )

    with open(out_path_md, 'w', encoding='utf-8') as f:
        temp_table = None
        table_name = []
        prev_page_text = ""

        for page in tqdm(pdf.pages):
            page_num = page.page_number
            
            # 合并曲线和边缘
            all_lines = page.curves + page.edges
            
            # 确保垂直线和水平线都至少有两个值
            vertical_lines = sorted(set([line['x0'] for line in all_lines] + [line['x1'] for line in all_lines]))
            horizontal_lines = sorted(set([line['top'] for line in all_lines] + [line['bottom'] for line in all_lines]))
            
            # 如果线条数量不足两个，添加页面边界
            if len(vertical_lines) < 2:
                vertical_lines = [0, page.width] + vertical_lines
            if len(horizontal_lines) < 2:
                horizontal_lines = [0, page.height] + horizontal_lines
            
            tables = page.find_tables(
                table_settings={
                    "vertical_strategy": "explicit",
                    "horizontal_strategy": "explicit",
                    "explicit_vertical_lines": vertical_lines,
                    "explicit_horizontal_lines": horizontal_lines,
                }
            )
            bboxes = [table.bbox for table in tables]

            page_text = page.filter(not_within_bboxes).extract_text()
            page_text = "\n".join(page_text.split("\n")[1:-1])
            if prev_page_text:
                last_line = prev_page_text.split("\n")[-1]
                first_line = page_text.split("\n")[0]

                if not is_sentence_end(last_line) and is_sentence_end(last_line + first_line):
                    f.write(f"# Page {page_num - 1} - Text\n\n")
                    f.write(prev_page_text + first_line)
                    f.write("\n\n")
                    page_text = "\n".join(page_text.split("\n")[1:])
                else:
                    f.write(f"# Page {page_num - 1} - Text\n\n")
                    f.write(prev_page_text)
                    f.write("\n\n")

            if page_text.strip():
                prev_page_text = page_text
            else:
                prev_page_text = ""
# 保存页面中的图片
            images = page.images
            for image in images:
                image_name = f"page_{page_num}_image_{images.index(image)}.png"
                image_path = os.path.join(image_dir, image_name)
                with open(image_path, "wb") as image_file:
                    image_file.write(image["stream"].get_data())
                       # 将图片信息写入Markdown文档
                #将图片的ocr内容写入Markdown
                '''
                    with open(image_path, "rb") as file:
                        files = {"file": file}
                        response = requests.post(url_f, files=files)
                    
                    outs = response.json()["detection_result"]
                    all_text += '图片识别内容'
                    all_text += outs
                    all_text += "\n"
                '''
                
                #f.write(f"# Page {page_num} - Image {images.index(image)}\n\n")
                #f.write(f"![Page {page_num} - Image {images.index(image)}]({image_path})\n\n")

            for table_id, table in enumerate(page.extract_tables(), start=1):
                if temp_table is not None:
                    if page.bbox[3] - tables[0].bbox[1] + 40 >= page.chars[0].get('y1'):
                        df = pd.DataFrame(table)
                        temp_table = pd.concat([temp_table, df], axis=0)
                        table_name.append(f"{page_num}-{table_id}")
                        if page.chars[-1].get('y0') < page.bbox[3] - tables[0].bbox[3]:
                            if is_valid_table(temp_table):
                                f.write(f"# Page {page_num} - Table {table_id}\n\n")
                                f.write(temp_table.to_markdown(index=False))
                                f.write("\n\n")
                            temp_table = None
                            table_name = []
                        else:
                            break
                    else:
                        if is_valid_table(temp_table):
                            f.write(f"# Page {page_num} - Table {table_name[0].split('-')[1]}\n\n")
                            f.write(temp_table.to_markdown(index=False))
                            f.write("\n\n")
                        temp_table = None
                        table_name = []
                        if page.chars[-1].get('y0') < page.bbox[3] - tables[table_id - 1].bbox[3]:
                            df = pd.DataFrame(table)
                            if is_valid_table(df):
                                f.write(f"# Page {page_num} - Table {table_id}\n\n")
                                f.write(df.to_markdown(index=False))
                                f.write("\n\n")
                        else:
                            temp_table = pd.DataFrame(table)
                            table_name.append(f"{page_num}-{table_id}")
                else:
                    if page.chars[-1].get('y0') < page.bbox[3] - tables[table_id - 1].bbox[3]:
                        df = pd.DataFrame(table)
                        if is_valid_table(df):
                            f.write(f"# Page {page_num} - Table {table_id}\n\n")
                            f.write(df.to_markdown(index=False))
                            f.write("\n\n")
                    else:
                        temp_table = pd.DataFrame(table)
                        table_name.append(f"{page_num}-{table_id}")
    pdf.close()


#常规md拆分方式-------固定头拆解

headers_to_split_on = [
    ("#", "Header 1"),
    ("##", "Header 2"),
    ("###", "Header 3"),
    ("####", "Header 4")
]

#递归拆解方式-----根据各种不同的头的拆解，例如**----**，***----***，****----****等等
separators=["#####","######","#######","####", "###","\n**"]
    

    # 定义拆分器
markdown_splitter = MarkdownHeaderTextSplitter(headers_to_split_on=headers_to_split_on)


def split_text_preserving_tables(text, max_token_length=3000):
    paragraphs = text.split('\n')
    chunks = []
    current_chunk = []
    current_length = 0
    
    for para in paragraphs:
        para_length = len(para)
        if current_length + para_length > max_token_length:
            chunks.append("\n".join(current_chunk))
            current_chunk = [para]
            current_length = para_length
        else:
            current_chunk.append(para)
            current_length += para_length
    
    if current_chunk:
        chunks.append("\n".join(current_chunk))
    
    return chunks


def process_doc_file(doc_file, image_output_dir, markdown_directory):
    file_extension = os.path.splitext(doc_file)[1].lower()
    md_header_splits = []
    
    if file_extension == ".docx":
        try:
            # 获取文档的基文件名（不带路径和扩展名）
            base_name = os.path.splitext(os.path.basename(doc_file))[0]
            
            # 创建该文档的专属媒体目录
            media_dir = os.path.join(image_output_dir, base_name)
            os.makedirs(media_dir, exist_ok=True)
            
            # 将 .docx 文件转换为 Markdown，提取的媒体文件保存在 media_dir 中
            # If pandoc is unavailable or conversion fails, fallback to python-docx
            try:
                markdown_text = pypandoc.convert_file(doc_file, 'markdown', extra_args=['--extract-media=' + media_dir])
            except Exception as e:
                # 减少警告信息的详细程度
                logging.info(f"Pandoc not available, using python-docx as fallback.")
                try:
                    docu = Document2(doc_file)
                    markdown_text = ""
                    for p in docu.paragraphs:
                        text = p.text.strip()
                        markdown_text += (text + "\n") if text else "\n"
                except Exception as e2:
                    logging.error(f"python-docx fallback also failed for {doc_file}: {e2}.")
                    markdown_text = ""
            
            c = 1
            img_md_pattern = r"!\[[^\]]*\]\(([^)]+)\)"
            base_path = os.path.join(media_dir, "media")

            all_text = ""
            img_found = 0
            for i in markdown_text.split("\n"):
                handled_image = False
                if "media/" in i:
                    matches = re.findall(img_md_pattern, i)
                    img_path = None
                    if matches:
                        rel = matches[0]
                        if rel.startswith("media/"):
                            img_path = os.path.join(media_dir, rel.replace('/', os.sep))
                            handled_image = True
                    if not handled_image:
                        m = re.search(r"media/(image\d+\.(?:png|jpe?g|gif|bmp))", i, flags=re.IGNORECASE)
                        if m:
                            img_path = os.path.join(base_path, m.group(1))
                            handled_image = True

                    if handled_image and img_path and ENABLE_OCR_IMAGES:
                        try:
                            logging.info(f"[OCR] request {url_f} file={img_path}")
                            with open(img_path, "rb") as file:
                                files = {"file": file}
                                response = requests.post(url_f, files=files, timeout=OCR_TIMEOUT)
                            response.raise_for_status()
                            outs = response.json().get("detection_result", "")
                            if outs:
                                all_text += "图片识别内容\n" + outs + "\n"
                                img_found += 1
                            logging.info(f"[OCR] success file={img_path} len={len(outs)}")
                        except Exception:
                            logging.exception(f"[OCR] failed file={img_path}")

                    if "height=" not in i:
                        c = 0
                elif c != 0:
                    all_text += i + "\n"
                else:
                    c = 1

            if ENABLE_OCR_IMAGES and img_found == 0:
                try:
                    out_media = os.path.join(media_dir, "media")
                    os.makedirs(out_media, exist_ok=True)
                    with zipfile.ZipFile(doc_file) as zf:
                        for name in zf.namelist():
                            nl = name.lower()
                            if nl.startswith("word/media/") and (nl.endswith(".png") or nl.endswith(".jpg") or nl.endswith(".jpeg") or nl.endswith(".bmp") or nl.endswith(".gif")):
                                dest = os.path.join(out_media, os.path.basename(name))
                                if not os.path.exists(dest):
                                    with zf.open(name) as src, open(dest, "wb") as dst:
                                        dst.write(src.read())
                                try:
                                    logging.info(f"[OCR] request {url_f} file={dest}")
                                    with open(dest, "rb") as file:
                                        files = {"file": file}
                                        response = requests.post(url_f, files=files, timeout=OCR_TIMEOUT)
                                    response.raise_for_status()
                                    outs = response.json().get("detection_result", "")
                                    if outs:
                                        all_text += "图片识别内容\n" + outs + "\n"
                                    logging.info(f"[OCR] success file={dest} len={len(outs)}")
                                except Exception:
                                    logging.exception(f"[OCR] failed file={dest}")
                except Exception:
                    logging.exception(f"[DOCX] media unzip failed file={doc_file}")
            
            #docu = Document2(doc_file)
            """
            markdown_text = ""
            for i in docu.paragraphs:
                if i.text == "":
                    markdown_text += "\n"
                else:
                    markdown_text += i.text
            """
            if len(all_text.replace(" ","").replace("\n","")) == 0:
                raise ValueError(f"The file {os.path.basename(doc_file)} does not contain valid content.")
            # 生成 Markdown 文件的路径
            markdown_file = os.path.join(markdown_directory, os.path.splitext(os.path.basename(doc_file))[0] + ".md")
            
            # 将转换后的 Markdown 文本写入文件
            with open(markdown_file, 'w', encoding='utf-8') as file:
                file.write(all_text)
            
            # 加载 Markdown 文件
            loader = TextLoader(markdown_file)
            document = loader.load()[0]
            
            # 拆分 Markdown 内容
            chunk_logger.info("正在分块 DOCX 文档...")
            text_splitter = RecursiveCharacterTextSplitter(separators=["#####", "######", "#######", "####", "###", "\n**"], chunk_size=200, chunk_overlap=30)
            md_header_splits = text_splitter.create_documents([all_text])
            
            if len(md_header_splits) == 0:
                # 如果没有拆分,将整个文档内容作为一个拆分
                md_header_splits = [document]
            else:
                for split in md_header_splits:
                    # 获取每个分块的元数据
                    metadata = split.metadata
                    
                    # 从分块内容中提取所有以#开头的行
                    lines = split.page_content.strip().split("\n")
                    title_lines = []
                    for line in lines:
                        if line.startswith("#"):
                            title_lines.append(line.strip("#").strip())
                    
                    # 将提取出的标题行作为 docuname
                    if title_lines:
                        metadata["docuname"] = "\n".join(title_lines)
                    
                    # 获取元数据中最后一级的标题
                    last_header = None
                    headers_to_split_on = ["#####", "######", "#######", "####", "###"]
                    for header_level in reversed(headers_to_split_on):
                        if header_level in metadata:
                            last_header = metadata[header_level]
                            break
                    
                    # 将最后一级标题添加到拆分部分的内容前面
                    if last_header:
                        split.page_content = f"## {last_header}\n\n{split.page_content}"
                    split.metadata["file_path"] = doc_file
            
            # 如果只有一个拆分且内容长度超过500 token，按段落进行切割并重新拆分
            if len(md_header_splits) == 1 and len(md_header_splits[0].page_content.split()) > 500:
                original_content = md_header_splits[0].page_content
                split_contents = split_text_preserving_tables(original_content)
                
                # 处理每个分割后的文本块，移除不必要的换行符
                processed_splits = []
                for chunk in split_contents:
                    # 保留表格的结构，但移除其他不必要的换行符
                    lines = chunk.split('\n')
                    processed_lines = []
                    in_table = False
                    for line in lines:
                        if line.strip().startswith('|') and line.strip().endswith('|'):
                            # 这是表格的一行，保留原样
                            processed_lines.append(line)
                            in_table = True
                        elif in_table and line.strip() == '':
                            # 表格结束
                            in_table = False
                            processed_lines.append(line)
                        elif not in_table:
                            # 非表格内容，移除换行符
                            processed_lines.append(line.strip())
                    
                    # 将处理后的行重新组合成一个文本块
                    processed_chunk = ' '.join(processed_lines).strip()
                    processed_splits.append(processed_chunk)
                
                # 创建新的Document对象
                md_header_splits = [Document(page_content=chunk, metadata={"file_path": doc_file}) for chunk in processed_splits]
        
        except Exception as e:
            logging.error(f"处理 {doc_file} 发生错误: {e}. 跳过该文件。")
    
    return md_header_splits

def process_md_file(md_file,markdown_directory):
    loader = TextLoader(md_file)
    document = loader.load()[0]
    base_name = os.path.basename(md_file)
    if len(document.page_content.replace(" ","").replace("\n","")) == 0:
        if os.path.exists(md_file):
            os.remove(md_file)
            print(f"文件 '{md_file}' 已被删除。")
        else:
            print(f"文件 '{md_file}' 不存在。")
        raise ValueError(f"The file {base_name} does not contain valid content.")
        
    markdown_file = os.path.join(markdown_directory, f"{os.path.splitext(os.path.basename(md_file))[0]}.md")
    with open(markdown_file, 'w', encoding='utf-8') as file:
        file.write(document.page_content)
    if '####' in document.page_content and '#' in document.page_content[0]:
        md_header_splits = []
        for chunk in document.page_content.split("####")[1:]:
            lines = [line.strip() for line in chunk.split("\n") if line.strip()]
            if 'http' in lines[-1]:
                md_header_splits.append(Document(page_content="\n".join(lines[:-1]), metadata={"file_path": md_file,"file_url":lines[-1],"isQA":1}))
            else:
                md_header_splits.append(Document(page_content="\n".join(lines), metadata={"file_path": md_file,"file_url":'-',"isQA":1}))
        
        # 分块日志（QA格式的Markdown）
        try:
            chunk_logger.info(f"Markdown(QA) 分块: file={os.path.basename(md_file)} chunks={len(md_header_splits)}")
            if md_header_splits:
                preview = md_header_splits[0].page_content[:200].replace("\n", " ")
                chunk_logger.info(f"Markdown(QA) 首块预览: {preview}")
        except Exception:
            pass
        return md_header_splits

    else:
        # 拆分 Markdown 内容
        chunk_logger.info("正在分块 Markdown 文档...")
        text_splitter = RecursiveCharacterTextSplitter(separators=["#####","######","#######","####", "###","\n**"], chunk_size=300, chunk_overlap=30)
        md_header_splits = text_splitter.create_documents([document.page_content])
        
        if len(md_header_splits) == 0:
            # 如果没有拆分,将整个文档内容作为一个拆分
            md_header_splits = [document]
        else:
            for split in md_header_splits:
                # 获取每个分块的元数据
                metadata = split.metadata
                
                # 从分块内容中提取所有以#开头的行
                lines = split.page_content.strip().split("\n")
                title_lines = []
                for line in lines:
                    if line.startswith("#"):
                        title_lines.append(line.strip("#").strip())
                
                # 将提取出的标题行作为docuname
                if title_lines:
                    metadata["docuname"] = "\n".join(title_lines)
                
                # 获取元数据中最后一级的标题
                last_header = None
                headers_to_split_on = ["#####", "######", "#######", "####", "###"]
                for header_level in reversed(headers_to_split_on):
                    if header_level in metadata:
                        last_header = metadata[header_level]
                        break
                
                # 将最后一级标题添加到拆分部分的内容前面
                if last_header:
                    split.page_content = f"## {last_header}\n\n{split.page_content}"
                split.metadata["file_path"] = md_file

                # # 移除文本块中的换行符
                # paragraphs = split.page_content.split('\n\n')
                # cleaned_paragraphs = [p.replace('\n', ' ').strip() for p in paragraphs]
                # split.page_content = '\n\n'.join(cleaned_paragraphs)

        # 分块日志（普通Markdown）
        try:
            chunk_logger.info(f"Markdown 分块: file={os.path.basename(md_file)} chunks={len(md_header_splits)}")
            if md_header_splits:
                preview = md_header_splits[0].page_content[:200].replace("\n", " ")
                chunk_logger.info(f"Markdown 首块预览: {preview}")
        except Exception:
            pass
    return md_header_splits


def process_txt_file(txt_file, markdown_directory):
    with open(txt_file, 'r', encoding='utf-8') as file:
        text = file.read()
        
    base_name = os.path.basename(txt_file)
    if len(text.replace(" ","").replace("\n","")) == 0:
        
        if os.path.exists(txt_file):
            os.remove(txt_file)
            print(f"文件 '{txt_file}' 已被删除。")
        else:
            print(f"文件 '{txt_file}' 不存在。")
        raise ValueError(f"The file {base_name} does not contain valid content.")
    
    # 拆分 Markdown 内容
    chunk_logger.info("正在分块 TXT 文档...")
    text_splitter = RecursiveCharacterTextSplitter(separators=["#####","######","#######","####", "###","\n**","\n\n",""], chunk_size=2000, chunk_overlap=200)
    md_header_splits = text_splitter.create_documents([text])
    
    if len(md_header_splits) == 0:
        # 如果没有拆分,将整个文档内容作为一个拆分
        md_header_splits = [Document(page_content=text)]
    else:
        # 创建 Markdown 文件路径
        markdown_file = os.path.join(markdown_directory, f"{os.path.splitext(os.path.basename(txt_file))[0]}.md")
        
        with open(markdown_file, 'w', encoding='utf-8') as file:
            for index, split in enumerate(tqdm(md_header_splits, desc="Processing splits")):
                # 获取每个分块的元数据
                metadata = split.metadata
                
                # 从分块内容中提取所有以#开头的行
                lines = split.page_content.strip().split("\n")
                title_lines = [line.strip("#").strip() for line in lines if line.startswith("#")]
                
                # 将提取出的标题行作为docuname
                if title_lines:
                    metadata["docuname"] = "\n".join(title_lines)
                
                # 获取元数据中最后一级的标题
                last_header = None
                headers_to_split_on = ["#####", "######", "#######", "####", "###"]
                for header_level in reversed(headers_to_split_on):
                    if header_level in metadata:
                        last_header = metadata[header_level]
                        break
                
                # 将最后一级标题添加到拆分部分的内容前面
                if last_header:
                    split.page_content = f"## {last_header}\n\n{split.page_content}"
                
                split.metadata["file_path"] = txt_file
                
                # 移除文本块中的换行符
                paragraphs = split.page_content.split('\n\n')
                cleaned_paragraphs = [p.replace('\n', ' ').strip() for p in paragraphs]
                cleaned_content = '\n\n'.join(cleaned_paragraphs)
                
                # 将处理后的分块内容写入文件
                file.write(cleaned_content)
                file.write("\n\n---\n\n")  # 添加分块分隔标志
    
    # 分块日志（TXT）
    try:
        chunk_logger.info(f"TXT 分块: file={os.path.basename(txt_file)} chunks={len(md_header_splits)}")
        if md_header_splits:
            preview = md_header_splits[0].page_content[:200].replace("\n", " ")
            chunk_logger.info(f"TXT 首块预览: {preview}")
    except Exception:
        pass
    return md_header_splits


def pdf_to_markdown(url, pdf_file_path, markdown_file_path, extract_images=False):
    """
    Convert a PDF file to Markdown format by sending a request to the specified URL.

    :param url: The URL to send the POST request to.
    :param pdf_file_path: The file path of the PDF to be converted.
    :param markdown_file_path: The file path where the resulting Markdown will be saved.
    :param extract_images: Optional parameter to extract images from the PDF.
    """
    try:
        # Open and read the PDF file
        #images = convert_from_path('/root/autodl-tmp/project_knowledge/a1.pdf')
        

        """
        with open(pdf_file_path, 'rb') as pdf_file:
            pdf_content = pdf_file.read()

        # Prepare the files and parameters for the request
        files = {'pdf_files': (os.path.basename(pdf_file_path), pdf_content, 'application/pdf')}
        params = {'extract_images': extract_images}

        # Send the POST request
        response = requests.post(url, files=files, params=params)
        response.raise_for_status()  # Raise an HTTPError if the HTTP request returned an unsuccessful status code

        # Get the response content as JSON
        response_json = response.json()[0]['markdown']

        # Convert the JSON response to Markdown format
        markdown_content = json_to_markdown(response_json)

        # Write the formatted Markdown content into the file
        with open(markdown_file_path, 'w', encoding='utf-8') as markdown_file:
            markdown_file.write(markdown_content)
        """
        data = {
            "file_type": "page",
            "resized_shape": 768,
            "embed_sep": " $,$ ",
            "isolated_sep": "$$\n, \n$$"
        }
        images = convert_from_path(pdf_file_path)

        all_text = ''
        for i in images:
            image = Image.fromarray(np.array(i))
            image_byte_array = io.BytesIO()
            image.save(image_byte_array, format='JPEG')
            image_byte_array = image_byte_array.getvalue()
            
            files = {'image': ('filename.jpg', image_byte_array, 'image/jpeg')}
            logging.info(f"[pix2text] request {url}")
            r = requests.post(url, data=data, files=files)
            
            all_text += r.json()['results']
            logging.info("[pix2text] success")
        
        with open(markdown_file_path, 'w', encoding='utf-8') as markdown_file:
            markdown_file.write(all_text)
        
        print(f"Response content has been written to {markdown_file_path}")

    except requests.exceptions.RequestException as e:
        print(f"An error occurred while making the request: {e}")
    except json.JSONDecodeError:
        print("An error occurred while parsing the response JSON.")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")

def json_to_markdown(data, level=1):
    """
    Convert JSON data to Markdown format.

    :param data: The JSON data to convert.
    :param level: The current heading level.
    :return: The formatted Markdown content.
    """
    md_content = ""
    if isinstance(data, dict):
        for key, value in data.items():
            md_content += f"{'#' * level} {key}\n\n"
            md_content += json_to_markdown(value, level + 1)
    elif isinstance(data, list):
        for item in data:
            md_content += json_to_markdown(item, level)
    else:
        md_content += f"{data}\n\n"
    return md_content

url = URL_PIX2TEXT


def process_pdf(pdf_path, image_dir, md_dir, chunk_size=1000):
    try:
        print(f"开始处理文档: {pdf_path}")
        chunk_logger.info("正在解析 PDF 并分块...")
        pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
        out_path_md = os.path.join(md_dir, f"{pdf_name}.md") 
        image_dir_file = os.path.join(image_dir, pdf_name)
        os.makedirs(image_dir_file, exist_ok=True)  # 为每个文档创建单独的图片文件夹
        
        try:
            extract_all_table_and_textand_image(pdf_path, out_path_md, image_dir_file)
            print("处理pdf完成-----")
            use_traditional_parsing = False
        except Exception as e:
            print(f"使用 extract_all_table_and_textand_image 解析文档出错: {e}")
            use_traditional_parsing = True
            if ENABLE_PDF_PIX2TEXT:
                pdf_to_markdown(url, pdf_path, out_path_md, extract_images=False)
            else:
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        try:
                            text += page.extract_text() or ""
                        except Exception:
                            pass
                with open(out_path_md, 'w', encoding='utf-8') as file:
                    file.write(text)
            

        # 检查生成的 Markdown 文件是否为空
        if not os.path.exists(out_path_md):
            base_name = os.path.basename(pdf_path)
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
                print(f"源文件 '{pdf_path}' 已被删除。")
            raise ValueError(f"未能生成Markdown文件，源文件 {base_name} 可能不包含有效内容。")
        
        loader = TextLoader(out_path_md)
        document = loader.load()[0]
        base_name = os.path.basename(pdf_path)
        if len(document.page_content.replace(" ","").replace("\n","")) == 0:
            if not use_traditional_parsing:
                if ENABLE_PDF_PIX2TEXT:
                    pdf_to_markdown(url, pdf_path, out_path_md, extract_images=False)
                else:
                    with open(pdf_path, 'rb') as file:
                        pdf_reader = PdfReader(file)
                        text = ""
                        for page in pdf_reader.pages:
                            try:
                                text += page.extract_text() or ""
                            except Exception:
                                pass
                    with open(out_path_md, 'w', encoding='utf-8') as file:
                        file.write(text)
                use_traditional_parsing = True
                loader = TextLoader(out_path_md)
                document = loader.load()[0]
            if len(document.page_content.replace(" ","").replace("\n","")) == 0:
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
                    if os.path.exists(out_path_md):
                        os.remove(out_path_md)
                    print(f"文件 '{pdf_path}' 已被删除。")
                else:
                    print(f"文件 '{pdf_path}' 不存在。")
                raise ValueError(f"The file {base_name} does not contain valid content.")

        if use_traditional_parsing:
            # 使用传统解析方法,按照指定的 chunk 大小分割文本
            text_splits = []
            text = document.page_content
            for i in range(0, len(text), chunk_size):
                chunk = text[i:i+chunk_size]
                metadata = {"document_name": chunk[:20], "file_path": pdf_path}
                text_splits.append(Document(page_content=chunk, metadata=metadata))
        else:
            # 使用原来的分割方式
            chunk_logger.info("正在按标题分割 PDF Markdown...")
            md_header_splits = markdown_splitter.split_text(document.page_content)
            
            if len(md_header_splits) > 0:
                # 判断文档名长度
                if len(pdf_name) >= 12:  # 假设6个中文字符等于12个英文字符
                    document_title = pdf_name
                else:
                    # 查找 "# Page 1 - Text" 对应的文本内容
                    page1_text = ""
                    for split in md_header_splits:
                        if "# Page 1 - Text" in split.page_content:
                            page1_text = split.page_content.split("# Page 1 - Text")[-1].strip()
                            break
                    
                    if len(page1_text) > 300:
                        document_title = page1_text[:150]  # 取前50个中文字符作为标题
                    else:
                        document_title = page1_text
                
                document.metadata["document_name"] = document_title
                
                # 给拆分后的每个元素添加文档名和文件路径的 metadata
                for split in md_header_splits:
                    split.metadata["document_name"] = document_title 
                    split.metadata["file_path"] = pdf_path
            
            text_splits = md_header_splits
        
        try:
            chunk_logger.info(f"PDF 分块: file={os.path.basename(pdf_path)} chunks={len(text_splits)}")
            if text_splits:
                preview = text_splits[0].page_content[:200].replace("\n", " ")
                chunk_logger.info(f"PDF 首块预览: {preview}")
        except Exception:
            pass
        print(f"文档处理完成: {pdf_path}")
        return text_splits
    
    except Exception as e:
        print(f"处理文档出错: {pdf_path}, 错误信息: {e}")
        return []
    
